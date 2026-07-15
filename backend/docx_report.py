"""
backend/docx_report.py
-------------------------
Generates a clean, professional downloadable Word (.docx) report of the
student's profile, skill-gap analysis, roadmap, courses, certifications and
progress summary using python-docx. Returns raw bytes so Streamlit's
`st.download_button` can serve it directly (no disk writes required).
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from backend.logger_setup import get_logger
from backend.recommendations import get_courses_for_domain, get_certifications_for_domain

logger = get_logger(__name__)

PRIMARY_RGB = RGBColor(0x7C, 0x5C, 0xFF)
DARK_RGB = RGBColor(0x1B, 0x1E, 0x33)
GREY_RGB = RGBColor(0x5B, 0x60, 0x79)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY_RGB if level == 1 else DARK_RGB


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 5"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)


def build_docx_report(
    profile: dict[str, Any],
    roadmap: dict[str, Any],
    completed_weeks: int = 0,
    readiness_percent: int = 0,
) -> bytes:
    """Build a full DOCX career report and return it as raw bytes."""
    doc = Document()

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LearnMate AI — Career Learning Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PRIMARY_RGB

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = GREY_RGB

    doc.add_paragraph()

    # ---- Student profile ----
    _heading(doc, "Student Profile", level=1)
    _add_table(
        doc,
        ["Field", "Details"],
        [
            ["Name", profile.get("name", "-")],
            ["Career Goal", profile.get("career_goal", "-")],
            ["Current Level", profile.get("current_level", "-")],
            ["Preferred Domain", profile.get("preferred_domain", "-")],
            ["Learning Preference", profile.get("learning_preference", "-")],
            ["Study Hours / Week", str(profile.get("study_hours_per_week", "-"))],
            ["Existing Skills", ", ".join(profile.get("existing_skills", []) or []) or "None listed"],
        ],
    )
    doc.add_paragraph()

    # ---- Progress summary ----
    _heading(doc, "Progress Summary", level=1)
    total_weeks = len(roadmap.get("weekly_milestones", []))
    p = doc.add_paragraph()
    p.add_run(f"Weeks completed: {completed_weeks} / {total_weeks}\n").bold = False
    p.add_run(f"Overall skill readiness: {readiness_percent}%\n")
    p.add_run(f"Estimated roadmap duration: {roadmap.get('estimated_duration_weeks', '-')} weeks")
    doc.add_paragraph()

    # ---- Roadmap overview ----
    _heading(doc, "Roadmap Overview", level=1)
    doc.add_paragraph(roadmap.get("summary", "No summary available."))

    _heading(doc, "Weekly Milestones", level=2)
    for m in roadmap.get("weekly_milestones", []):
        wk_p = doc.add_paragraph()
        wk_p.add_run(f"Week {m.get('week', '-')}: {m.get('title', '')}").bold = True
        if m.get("description"):
            doc.add_paragraph(m["description"])
        for topic in m.get("key_skills", []) or m.get("topics", []) or []:
            doc.add_paragraph(f"• {topic}", style="List Bullet")

    if roadmap.get("capstone_project"):
        _heading(doc, "Capstone Project", level=2)
        doc.add_paragraph(roadmap["capstone_project"])

    # ---- Skill gap ----
    gaps = roadmap.get("skill_gap_analysis", [])
    if gaps:
        _heading(doc, "Skill Gap Analysis", level=1)
        _add_table(
            doc,
            ["Skill", "Current Level", "Required Level", "Priority"],
            [
                [
                    g.get("skill", "-"),
                    str(g.get("current_level", "-")),
                    str(g.get("required_level", "-")),
                    g.get("priority", "-"),
                ]
                for g in gaps
            ],
        )
        doc.add_paragraph()

    # ---- Courses ----
    domain = profile.get("preferred_domain", "")
    _heading(doc, "Recommended Courses", level=1)
    _add_table(
        doc,
        ["Course", "Provider", "Duration", "Difficulty", "Pricing", "Link"],
        [[c.name, c.provider, c.duration, c.difficulty, c.pricing, c.link] for c in get_courses_for_domain(domain)],
    )
    doc.add_paragraph()

    # ---- Certifications ----
    _heading(doc, "Recommended Certifications", level=1)
    _add_table(
        doc,
        ["Certification", "Provider", "Difficulty", "Exam Cost", "Link"],
        [[c.name, c.provider, c.difficulty, c.exam_cost, c.link] for c in get_certifications_for_domain(domain)],
    )

    # ---- Footer ----
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer.add_run("Generated by LearnMate AI — powered by IBM watsonx.ai (Granite models)")
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = GREY_RGB

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
